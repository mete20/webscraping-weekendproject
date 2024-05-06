import requests
from bs4 import BeautifulSoup
from docx import Document
import re
from datetime import datetime

# URL of the webpage you want to scrape
url = 'https://www.mhp.org.tr/htmldocs/genel_baskan/1565/konusmalari/Devlet_Bahceli_2022_yili_konusmalari.html'

# Words to search for in the link text
search_words = ['Grup']

# Send a GET request to the webpage
response = requests.get(url)

# Parse the HTML content
soup = BeautifulSoup(response.text, 'html.parser')

# Find all <a> tags with class "mhp_table_link"
links = soup.find_all('a', class_='mhp_table_link')

# Filter links that contain the search words
filtered_links = [link['href'] for link in links if any(word in link.text for word in search_words)]

# Print the filtered links


def scrape_and_save(url):
    # Send a GET request to the URL and fetch the HTML content
    response = requests.get(url)
    html_code = response.content

    # Parse the HTML content using Beautiful Soup
    soup = BeautifulSoup(html_code, 'html.parser')

    # Find the first paragraph with class "govdeverdana"
    first_paragraph = soup.find('p', class_='baslikbasinmetin') 

    title = first_paragraph.get_text() if first_paragraph else "Untitled"       


    # Turkish month names to English month names mapping
    turkish_month_names = {
        "Ocak": "01",
        "Şubat": "02",
        "Mart": "03",
        "Nisan": "04",
        "Mayıs": "05",
        "Haziran": "06",
        "Temmuz": "07",
        "Ağustos": "08",
        "Eylül": "09",
        "Ekim": "10",
        "Kasım": "11",
        "Aralık": "12"
    }

    # Extract text content from the first paragraph
    filename = "scrap.docx"
    if first_paragraph:
        first_paragraph_text = first_paragraph.get_text()
        
        # Use regex to find the date pattern in the first paragraph text
        date_match = re.search(r'\d+\s+\w+\s+\d{4}', first_paragraph_text)
        #date_match = re.search(r'\d+\s+(Ocak|Şubat|Mart|Nisan|Mayıs|Haziran|Temmuz|Ağustos|Eylül|Ekim|Kasım|Aralık)\s+\d{4}', first_paragraph_text)
        
        date_str = date_match.group(0) if date_match else None
        # Use regex to extract the day, month, and year from the date string
        day, month_name, year = re.findall(r'\d+|\w+', date_str)

        # Convert the month name to the corresponding English month name
        month = turkish_month_names.get(month_name)

        # Create a datetime object with the extracted date
        date = datetime(int(year), int(month), int(day))

        if date:
            # if the date is not between 1 Jan 2022, 31 October 2022 return
            if date < datetime(2022, 1, 1) or date > datetime(2022, 10, 31):
                print("Date not in range")
                return None

        # Print the date in the desired format
        formatted_date = date.strftime('%d.%m.%Y')
        filename = 'MHP - ' + formatted_date + ".docx"
        
    # Find all paragraphs with class "govdeverdana"
    paragraphs = soup.find_all('p', class_='govdeverdana')


    # Extract text content from each paragraph
    content = [p.get_text() for p in paragraphs]

    # Join the content into a single string
    content_text = '\n'.join(content)

    # Create a new Word document
    doc = Document()

    doc.add_heading(title, level=1)

    # Add the scraped content to the document
    doc.add_paragraph(content_text)

    # Save the document to a file with the extracted date as the filename
    doc.save(filename)

    print("Document saved successfully as:", filename)



for link in filtered_links:
    url = "https://www.mhp.org.tr/" + link

    
