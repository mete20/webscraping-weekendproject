import requests
from bs4 import BeautifulSoup
from docx import Document
import re
from datetime import datetime


def scrape_and_print_urls(urls):
    href_links = []
    for url in urls:
        # Send a GET request to the URL and fetch the HTML content
        response = requests.get(url)
        html_code = response.content

        # Parse the HTML content using Beautiful Soup
        soup = BeautifulSoup(html_code, 'html.parser')

        links = soup.select('.box-content-header a')

         # Extract the href attribute value from each link
        #href_links = [link.get('href') for link in links]
        for link in links:
            href_links.append(link.get('href'))
    return href_links        

# Example list of URLs
urls = [
    'https://hdp.org.tr/tr/grup-konusmalari/6033/sayfa/2/',
    'https://hdp.org.tr/tr/grup-konusmalari/6033/sayfa/3/',
    
]

# Call the function to scrape and print content from the provided URLs
href_links = scrape_and_print_urls(urls)

turkish_month_names = {
        "Ocak": "01",
        "Şubat": "02",
        "Mart": "03",
        "Nisan": "04",
        "Mayıs": "05",
        "Haziran": "06",
        "Temmuz": "07",
        "Ağustos": "08",
        "Eylül": "09",
        "Ekim": "10",
        "Kasım": "11",
        "Aralık": "12"
    }

def scrape_content_and_save(url, count = 0):
    # Send a GET request to the URL and fetch the HTML content
    response = requests.get(url)
    html_code = response.content

    # Parse the HTML content using Beautiful Soup
    soup = BeautifulSoup(html_code, 'html.parser')

    # Find title of the page
    title = soup.find('div', class_='page-header').find('h1').text.strip()


    page_content_div = soup.find('div', class_='page-content')
    if not page_content_div:
        print("Page content div not found")
        return

    # Find the last <p> tag to extract the date
    last_p_tag = page_content_div.find_all('p')[-1]
    date_str = last_p_tag.text.strip()

    # Print the extracted date string for debugging
    #print("Extracted Date String:", date_str)

    # Use regex to extract day, month name, and year from the date string
    date_match = re.search(r'\d+\s+\w+\s+\d{4}', date_str)
    if not date_match:
        last_p_tag = page_content_div.find_all('p')[-2]
        date_str = last_p_tag.text.strip()
        date_match = re.search(r'\d+\s+\w+\s+\d{4}', date_str)
        if not date_match:
            print("Date not found in expected format")
            print(url)
            return

    # Extract day, month name, and year from the matched date string
    day, month_name, year = re.findall(r'\d+|\w+', date_match.group(0))

    # Convert the month name to the corresponding English month name
    month = turkish_month_names.get(month_name)
    if not month:
        print("Invalid month name:", month_name)
        return

    # Create a datetime object with the extracted date
    date = datetime(int(year), int(month), int(day))

    # Check if the date is within the desired range
    if date < datetime(2022, 1, 1) or date > datetime(2022, 10, 31):
        print("Date not in range" + str(date)) 
        return

    # Format the date for naming the document
    formatted_date = date.strftime("%d.%m.%Y")

    # Find all <p> tags to extract the content
    paragraphs = page_content_div.find_all('p')[:-1]  # Exclude the last <p> containing the date

    # Extract text content from each paragraph
    content = [p.get_text() for p in paragraphs]

    # Join the content into a single string
    content_text = '\n'.join(content)

    # Create a new Word document
    doc = Document()

    doc.add_heading(title, level=1)

    # Add the scraped content to the document
    doc.add_paragraph(content_text)

    # Save the document with the formatted date as the filename
    filename = f"{formatted_date}.docx"
    filename = 'HDP - ' + filename
    doc.save(filename)

    print("Document saved successfully as:", filename)
    return count + 1

main_url = 'https://hdp.org.tr'

count = 0
for link in href_links:
    full_link = main_url + link
    result = scrape_content_and_save(full_link)
    count += result if result else 0

print("Total documents saved:", count)
    


    





